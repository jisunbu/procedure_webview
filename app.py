# 필요한 라이브러리 임포트
from flask import Flask, render_template, request, jsonify, send_file
from markupsafe import Markup  # HTML 안전하게 렌더링
import os
from docx import Document  # Word 문서 처리
from docx.shared import Inches
from docx.oxml.shared import qn
from docx.table import _Cell
import pandas as pd  # Excel 처리
import PyPDF2  # PDF 처리
from werkzeug.utils import secure_filename
from datetime import datetime
import os.path
import base64
from io import BytesIO
from PIL import Image
from config import config

def create_app(config_name='default'):
    app = Flask(__name__)
    app.config.from_object(config[config_name])
    
    # 업로드 폴더 설정
    upload_folder = os.getenv('UPLOAD_FOLDER', 'documents/company_docs')
    app.config['UPLOAD_FOLDER'] = upload_folder
    
    # 절대 경로로 변환
    if not os.path.isabs(app.config['UPLOAD_FOLDER']):
        app.config['UPLOAD_FOLDER'] = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            app.config['UPLOAD_FOLDER']
        )
    
    print(f"Upload folder path: {app.config['UPLOAD_FOLDER']}")
    
    # 폴더 생성
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    
    return app

app = create_app(os.getenv('FLASK_ENV', 'production'))

# 문서 저장 경로와 허용된 파일 확장자 설정
UPLOAD_FOLDER = os.getenv('UPLOAD_FOLDER', 'documents/company_docs')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx', 'xls'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # 캐시 비활성화

# 파일 확장자 검사 함수
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# PDF 파일 읽기 함수
def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

# Word 파일 읽기 함수
def read_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

# Excel 파일 읽기 함수
def read_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_html(
        classes='table table-striped table-bordered table-hover',
        index=False,
        na_rep='-',
        float_format=lambda x: '{:.2f}'.format(x) if pd.notnull(x) else '-'
    )

# 메인 페이지 라우트
@app.route('/')
def index():
    try:
        # 폴더 경로 출력 (디버깅용)
        print(f"Current UPLOAD_FOLDER: {app.config['UPLOAD_FOLDER']}")
        
        # 폴더 존재 확인
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            print(f"Created upload folder: {app.config['UPLOAD_FOLDER']}")
        
        # 폴더 내용 출력 (디버깅용)
        print("Directory contents:")
        for root, dirs, files in os.walk(app.config['UPLOAD_FOLDER']):
            print(f"Root: {root}")
            print(f"Dirs: {dirs}")
            print(f"Files: {files}")
        
        files = []
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            if allowed_file(filename):
                files.append(filename)
        
        files.sort()
        print(f"Found {len(files)} files")
        
        return render_template('index.html', files=files)
    except Exception as e:
        print(f"Error in index route: {str(e)}")
        return render_template('index.html', 
                             files=[], 
                             error=f"Error: {str(e)}")

# 파일 보기 라우트
@app.route('/view/<filename>')
def view_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {filename}")

        file_type = filename.rsplit('.', 1)[1].lower()
        
        if file_type == 'pdf':
            return send_file(file_path, mimetype='application/pdf')
            
        elif file_type == 'docx':
            doc = Document(file_path)
            html_content = []
            
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    for paragraph in doc.paragraphs:
                        style = paragraph.style.name
                        text = paragraph.text
                        
                        if style.startswith('Heading'):
                            level = style[-1]
                            html_content.append(f'<h{level}>{text}</h{level}>')
                        else:
                            html_content.append(f'<p>{text}</p>')
                            
                elif element.tag.endswith('tbl'):
                    for table in doc.tables:
                        table_html = ['<table class="table table-bordered">']
                        for row in table.rows:
                            table_html.append('<tr>')
                            for cell in row.cells:
                                table_html.append(f'<td>{cell.text}</td>')
                            table_html.append('</tr>')
                        table_html.append('</table>')
                        html_content.append(''.join(table_html))
            
            content = Markup('\n'.join(html_content))
            return render_template('viewer.html', filename=filename, 
                                content=content, file_type=file_type)
            
        elif file_type in ['xlsx', 'xls']:
            content = Markup(read_excel(file_path))
            return render_template('viewer.html', filename=filename, 
                                content=content, file_type=file_type)
            
    except Exception as e:
        print(f"Error viewing file {filename}: {str(e)}")
        return f"Error viewing file: {str(e)}", 500

# 파일 다운로드 라우트
@app.route('/download/<filename>')
def download_file(filename):
    """파일 다운로드 처리"""
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {filename}")
            
        return send_file(file_path, as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Error downloading file {filename}: {str(e)}")
        return f"Error downloading file: {str(e)}", 500

@app.errorhandler(404)
def not_found_error(error):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return render_template('errors/500.html'), 500

# Flask 서버 실행 설정
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', 5000)))

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': '파일이 선택되지 않았습니다.'}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '파일이 선택되지 않았습니다.'}), 400
            
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            # 업로드 폴더 확인 및 생성
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            
            file.save(file_path)
            print(f"File uploaded successfully: {filename}")
            return jsonify({'success': True, 'filename': filename}), 200
            
        return jsonify({'error': '허용되지 않는 파일 형식입니다.'}), 400
    except Exception as e:
        print(f"Error uploading file: {str(e)}")
        return jsonify({'error': f'파일 업로드 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/test-upload', methods=['GET'])
def test_upload():
    try:
        test_file = os.path.join(app.config['UPLOAD_FOLDER'], 'test.pdf')
        with open(test_file, 'w') as f:
            f.write('Test PDF content')
        return jsonify({
            'success': True,
            'message': f'Test file created at {test_file}',
            'folder': app.config['UPLOAD_FOLDER'],
            'files': os.listdir(app.config['UPLOAD_FOLDER'])
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/debug-info')
def debug_info():
    try:
        upload_folder = app.config['UPLOAD_FOLDER']
        files = os.listdir(upload_folder) if os.path.exists(upload_folder) else []
        
        return jsonify({
            'upload_folder': upload_folder,
            'folder_exists': os.path.exists(upload_folder),
            'files': files,
            'env_vars': {
                'FLASK_ENV': os.getenv('FLASK_ENV'),
                'UPLOAD_FOLDER': os.getenv('UPLOAD_FOLDER')
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
